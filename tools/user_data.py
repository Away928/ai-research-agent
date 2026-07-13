"""
用户补充数据管理器 — 上传文件解析
==================================

支持格式：.md / .txt / .pdf / .docx
长文本策略：≤5000 字全文保留，>5000 字可选 AI 摘要 + 原文关键段落

统一接口：
    from tools.user_data import UserDataManager
    mgr = UserDataManager()
    entries = mgr.parse_uploads(uploaded_files, category="会议纪要")
    # 返回 [{"标题": "...", "日期": "...", "来源": "...", "内容": "...", "全文长度": ...}, ...]
"""

import re
from typing import Dict, List, Optional


# 超长文本阈值
_LONG_TEXT_THRESHOLD = 5000


class UserDataManager:
    """用户补充数据管理。"""

    name = "user_data"

    # 支持的数据类别
    CATEGORIES = ["会议纪要", "专家访谈", "调研笔记", "其他"]

    def parse_uploads(self, uploaded_files: list,
                      categories: list = None,
                      llm_client=None) -> List[Dict]:
        """解析用户上传的文件列表，返回结构化条目。

        Args:
            uploaded_files: 列表，每个元素可以是：
                - 文件路径 (str/Path)
                - (文件名, bytes内容) 元组（Streamlit file_uploader 格式）
                - dict: {"name": ..., "content": ...}（已预解析文本）
            categories: 对应的类别列表（与 files 一一对应），默认全部为"会议纪要"
            llm_client: 可选，LLM 客户端实例。若提供，超长文本会走 AI 摘要

        Returns:
            [{"标题": ..., "日期": ..., "来源": ..., "内容": ..., "_原始长度": ...}, ...]
        """
        if not uploaded_files:
            return []

        results = []
        for i, file_obj in enumerate(uploaded_files):
            # 优先从 dict 中取 category，其次从 categories 列表取，最后默认"会议纪要"
            if isinstance(file_obj, dict) and file_obj.get("category"):
                category = file_obj["category"]
            elif categories and i < len(categories):
                category = categories[i]
            else:
                category = "会议纪要"
            entry = self._parse_one(file_obj, category, llm_client)
            if entry:
                results.append(entry)

        return results

    @staticmethod
    def summarize_supplements(raw) -> dict:
        """将用户补充数据压缩为摘要，供后续阶段传入 AI prompt。"""
        if not raw:
            return {"总计": 0, "说明": "无用户补充数据"}
        result = {"总计": 0, "明细": {}}
        if isinstance(raw, list):
            for item in raw:
                cat = item.get("类别", "其他")
                if cat not in result["明细"]:
                    result["明细"][cat] = []
                content = item.get("内容", "")
                summary = content[:1500] if len(content) > 1500 else content
                result["明细"][cat].append({
                    "标题": item.get("标题", ""),
                    "日期": item.get("日期", ""),
                    "来源": item.get("来源", ""),
                    "摘要": summary,
                })
                result["总计"] += 1
        else:
            for cat in ["会议纪要", "专家访谈", "调研笔记", "其他"]:
                items = raw.get(cat, []) or []
                if not items:
                    continue
                result["明细"][cat] = []
                for item in items:
                    content = item.get("内容", "")
                    summary = content[:1500] if len(content) > 1500 else content
                    result["明细"][cat].append({
                        "标题": item.get("标题", ""),
                        "日期": item.get("日期", ""),
                        "来源": item.get("来源", ""),
                        "摘要": summary,
                    })
                    result["总计"] += len(items)
        if result["总计"] == 0:
            result["说明"] = "无用户补充数据"
        return result

    def _parse_one(self, file_obj, category: str, llm_client=None) -> Optional[Dict]:
        """解析单个文件。"""
        # 提取文件名和内容
        name, text = self._extract(file_obj)
        if not text or len(text.strip()) < 10:
            return None

        # 从文本中推测标题、日期、来源
        title, date, source, body = self._parse_frontmatter(text, name)

        # 长文本处理
        content = body
        original_len = len(body)
        if original_len > _LONG_TEXT_THRESHOLD and llm_client:
            summary = self._ai_summarize(body, llm_client)
            if summary:
                # 摘要放在原文前面，帮 AI 建立索引
                content = (
                    f"[文件摘要]\n{summary}\n\n"
                    f"[原文（前{_LONG_TEXT_THRESHOLD}字）]\n"
                    f"{body[:_LONG_TEXT_THRESHOLD]}"
                )

        return {
            "标题": title,
            "日期": date,
            "来源": source,
            "内容": content,
            "类别": category,
            "_原始长度": original_len,
            "_已压缩": original_len > _LONG_TEXT_THRESHOLD and llm_client is not None,
        }

    # ── 文件提取 ────────────────────────────────────────────────

    def _extract(self, file_obj):
        """从各种格式中提取 (文件名, 文本内容)。"""
        import io
        from pathlib import Path

        # Case 1: 文件路径
        if isinstance(file_obj, (str, Path)):
            p = Path(file_obj)
            if not p.exists():
                return None, None
            name = p.name
            text = self._read_by_suffix(name, p.read_bytes())
            return name, text

        # Case 2: (文件名, bytes) 元组（Streamlit UploadedFile）
        if isinstance(file_obj, tuple) and len(file_obj) == 2:
            name, data = file_obj
            text = self._read_by_suffix(name, data if isinstance(data, bytes) else data.encode())
            return name, text

        # Case 3: dict {"name": ..., "content": ...}
        if isinstance(file_obj, dict):
            name = file_obj.get("name", file_obj.get("标题", "未命名"))
            text = file_obj.get("content", file_obj.get("内容", ""))
            return name, text

        # Case 4: Streamlit UploadedFile 对象
        if hasattr(file_obj, "name") and hasattr(file_obj, "read"):
            data = file_obj.read()
            name = getattr(file_obj, "name", "未命名")
            text = self._read_by_suffix(name, data)
            return name, text

        return None, None

    def _read_by_suffix(self, filename: str, data: bytes) -> Optional[str]:
        """根据文件后缀读取文本。"""
        suffix = filename.lower().split(".")[-1] if "." in filename else ""

        if suffix in ("md", "txt", ""):
            try:
                return data.decode("utf-8")
            except UnicodeDecodeError:
                try:
                    return data.decode("gbk")
                except Exception:
                    return None

        if suffix == "pdf":
            return self._read_pdf(data)

        if suffix == "docx":
            return self._read_docx(data)

        # 未知格式：尝试当文本读
        try:
            return data.decode("utf-8")
        except Exception:
            return None

    def _read_pdf(self, data: bytes) -> Optional[str]:
        """从 PDF 二进制数据提取文字。"""
        import io

        # 尝试 pdfplumber
        try:
            import pdfplumber
            with pdfplumber.open(io.BytesIO(data)) as pdf:
                pages = []
                for page in pdf.pages:
                    t = page.extract_text()
                    if t:
                        pages.append(t)
                text = "\n".join(pages).strip()
                if len(text) > 50:
                    return text
        except Exception:
            pass

        # 降级：PyPDF2
        try:
            from PyPDF2 import PdfReader
            reader = PdfReader(io.BytesIO(data))
            pages = []
            for page in reader.pages:
                t = page.extract_text()
                if t:
                    pages.append(t)
            text = "\n".join(pages).strip()
            if len(text) > 50:
                return text
        except Exception:
            pass

        return None

    def _read_docx(self, data: bytes) -> Optional[str]:
        """从 DOCX 二进制数据提取文字。"""
        import io
        try:
            from docx import Document
            doc = Document(io.BytesIO(data))
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            # 也提取表格内容
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text for cell in row.cells if cell.text.strip())
                    if row_text.strip():
                        paragraphs.append(row_text)
            text = "\n".join(paragraphs)
            return text if len(text) > 20 else None
        except Exception:
            return None

    # ── 元信息提取 ──────────────────────────────────────────────

    def _parse_frontmatter(self, text: str, name: str):
        """从文本开头提取标题、日期、来源。"""
        lines = text.strip().split("\n")
        title = name.rsplit(".", 1)[0] if "." in name else name
        date = ""
        source = ""
        body_start = 0

        # 第一行：标题（# 开头）
        if lines and lines[0].startswith("#"):
            title = lines[0].lstrip("#").strip()
            body_start = 1

        # 第二行：日期
        if len(lines) > body_start:
            line = lines[body_start].strip()
            if line.startswith("##"):
                date = line.lstrip("#").strip()
                body_start += 1
            elif re.match(r"^\d{4}-\d{2}-\d{2}$", line):
                date = line
                body_start += 1

        # 第三行：来源
        if len(lines) > body_start:
            line = lines[body_start].strip()
            if "来源" in line or "Source" in line:
                source = re.split(r"[：:]", line, maxsplit=1)[-1].strip()
                body_start += 1

        # 从文件名推断日期（如果 frontmatter 没找到）
        if not date:
            date_match = re.search(r"(\d{4}-\d{2}-\d{2})", name)
            if date_match:
                date = date_match.group(1)

        body = "\n".join(lines[body_start:]).strip()
        return title, date, source, body

    # ── AI 摘要 ─────────────────────────────────────────────────

    def _ai_summarize(self, text: str, llm_client) -> Optional[str]:
        """对长文本生成结构化摘要，帮助 AI 建立索引。"""
        prompt = (
            "请对以下投研文档生成一个结构化摘要（200-300字），包含：\n"
            "1. 核心主题（1句话）\n"
            "2. 关键数据点（数字/百分比/趋势）\n"
            "3. 主要观点和判断\n"
            "4. 值得后续分析重点关注的信息\n\n"
            f"文档内容（前3000字）：\n{text[:3000]}"
        )
        try:
            system = "你是一个投研文档整理助手。只输出摘要，不要添加额外评论。"
            result = llm_client.ask(prompt, system=system, max_tokens=600)
            if result:
                return result.strip()
        except Exception:
            pass
        return None
